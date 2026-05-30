"""PKPM calculation report renderers — Markdown, Word, and PDF generation."""
from __future__ import annotations

import os
import re
import subprocess
import logging
from pathlib import Path
from typing import Any, Dict, List, Optional

_logger = logging.getLogger(__name__)


# ── Image helpers ────────────────────────────────────────────────────────


def _find_project_images(project_dir: Path) -> List[Path]:
    """Find structural images (BuildUp.BMP, PNG, JPG) in project directory."""
    images: List[Path] = []
    for pattern in ["BuildUp.BMP", "BuildUp.bmp", "*.png", "*.jpg", "*.jpeg"]:
        images.extend(project_dir.glob(pattern))
    # Deduplicate by stem, keep first found (BMP before PNG for same stem)
    seen: set[str] = set()
    result: List[Path] = []
    for p in images:
        key = p.stem.lower()
        if key not in seen:
            seen.add(key)
            result.append(p)
    return result


def _convert_bmp_to_png(bmp_path: Path) -> Optional[Path]:
    """Convert BMP to PNG for embedding. Returns PNG path."""
    png_path = bmp_path.with_suffix(".png")
    if png_path.exists():
        return png_path
    try:
        from PIL import Image
        img = Image.open(str(bmp_path))
        if img.mode == "P":
            img = img.convert("RGB")
        img.save(str(png_path), "PNG")
        return png_path
    except Exception as exc:
        _logger.warning("BMP→PNG conversion failed for %s: %s", bmp_path, exc)
        return None


# ── Markdown generation ─────────────────────────────────────────────────


def _kv_lines(text: str) -> List[str]:
    lines: List[str] = []
    for line in text.strip().split("\n"):
        line = line.strip()
        if line:
            lines.append(f"- {line}")
    return lines


def _generate_markdown(report: Dict[str, Any]) -> str:
    lines: List[str] = []
    lines.append("# PKPM SATWE 结构计算书")
    lines.append("")

    # Structural images
    images: List[str] = report.get("detailed", {}).get("images", [])
    if images:
        lines.append("## 结构平面图")
        lines.append("")
        for img in images:
            posix_path = Path(img).as_posix()
            lines.append(f"![结构平面图]({posix_path})")
            lines.append("")

    detailed = report.get("detailed", {})
    out_data = detailed.get("out_file_data", {})

    # Section 1: Design Parameters
    params = out_data.get("wmass_params", {})
    if params:
        lines.append("## 一、设计参数总信息")
        lines.append("")
        for key, label in [
            ("design_params", "总体信息"),
            ("wind_info_params", "风荷载参数"),
            ("earthquake_params", "地震参数"),
            ("material_params", "材料参数"),
            ("adjustment_params", "调整参数"),
            ("design_info_params", "设计信息"),
            ("load_combination_params", "荷载组合参数"),
        ]:
            section = params.get(key)
            if section:
                lines.append(f"### {label}")
                lines.append("")
                lines.extend(_kv_lines(section))
                lines.append("")

    # Section 2: Mass & Member Counts
    mass_table = out_data.get("mass_table", [])
    member_counts = out_data.get("member_counts", [])
    total_mass = out_data.get("total_mass", {})
    if mass_table or member_counts:
        lines.append("## 二、楼层质量与构件信息")
        lines.append("")
        if total_mass:
            for k, v in total_mass.items():
                label = {
                    "live_total": "活载总质量",
                    "dead_total": "恒载总质量",
                    "add_total": "附加总质量",
                    "struct_total": "结构总质量",
                }.get(k, k)
                lines.append(f"- **{label}**: {v} t")
            lines.append("")
        if mass_table:
            lines.append("### 各层质量")
            lines.append("")
            lines.append("| 层号 | 塔号 | 恒载质量(t) | 活载质量(t) | 质量比 |")
            lines.append("|------|------|------------|------------|--------|")
            for r in mass_table:
                lines.append(f"| {r['floor']} | {r['tower']} | {r['dead_mass']} | {r['live_mass']} | {r['mass_ratio']} |")
            lines.append("")
        if member_counts:
            lines.append("### 构件数量与层高")
            lines.append("")
            lines.append("| 层号 | 梁数 | 柱数 | 层高(m) | 累计高度(m) |")
            lines.append("|------|------|------|---------|------------|")
            for r in member_counts:
                lines.append(f"| {r['floor']} | {r['beam_count']} | {r['column_count']} | {r['height']} | {r['cumulative_height']} |")
            lines.append("")

    # Section 3: Unit Mass
    unit_mass = out_data.get("unit_mass", [])
    if unit_mass:
        lines.append("## 三、单位面积质量分布")
        lines.append("")
        lines.append("| 层号 | 塔号 | 单位面积质量(kg/m²) | 质量比 |")
        lines.append("|------|------|--------------------|--------|")
        for r in unit_mass:
            lines.append(f"| {r['floor']} | {r['tower']} | {r['unit_mass']} | {r['mass_ratio']} |")
        lines.append("")

    # Section 4: Wind Load
    wind_load = out_data.get("wind_load", [])
    if wind_load:
        lines.append("## 四、风荷载信息")
        lines.append("")
        lines.append("| 层号 | 风荷载X(kN) | 剪力X(kN) | 倾覆X(kN·m) | 风荷载Y(kN) | 剪力Y(kN) | 倾覆Y(kN·m) |")
        lines.append("|------|------------|----------|------------|------------|----------|------------|")
        for r in wind_load:
            lines.append(
                f"| {r['floor']} | {r['wind_x']} | {r['shear_x']} | {r['overturn_x']} | "
                f"{r['wind_y']} | {r['shear_y']} | {r['overturn_y']} |"
            )
        lines.append("")

    # Section 5: Modal Analysis
    modal_api = detailed.get("modal_analysis", [])
    wzq_periods = out_data.get("wzq_periods", [])
    modal_data = wzq_periods if wzq_periods else modal_api
    if modal_data:
        lines.append("## 五、模态分析")
        lines.append("")
        if wzq_periods:
            lines.append("| 振型号 | 周期(s) | 方向角(°) | 平动系数 | X平动 | Y平动 | 扭转系数 |")
            lines.append("|--------|---------|-----------|---------|-------|-------|---------|")
            for m in wzq_periods:
                lines.append(
                    f"| {m['mode']} | {m['period']} | {m['angle']} | {m['translation']} | "
                    f"{m['x_translation']} | {m['y_translation']} | {m['torsion']} |"
                )
        else:
            lines.append("| 振型号 | 周期(s) | 方向角(°) | 扭转系数 | X侧移 | Y侧移 |")
            lines.append("|--------|---------|-----------|----------|-------|-------|")
            for m in modal_api:
                lines.append(
                    f"| {m['index']} | {m['period_s']} | {m['angle']} | "
                    f"{m['torsion_ratio']} | {m['x_side']} | {m['y_side']} |"
                )
        lines.append("")

    # Section 6: Earthquake Response (CQC)
    wzq_base_shear = out_data.get("wzq_base_shear", [])
    if wzq_base_shear:
        lines.append("## 六、地震作用下的楼层反应（CQC）")
        lines.append("")
        for direction in ["X", "Y"]:
            dir_data = [r for r in wzq_base_shear if r["direction"] == direction]
            if dir_data:
                lines.append(f"### {direction}方向")
                lines.append("")
                lines.append("| 层号 | 塔号 | F(kN) | V(kN) |")
                lines.append("|------|------|-------|-------|")
                for r in dir_data:
                    lines.append(f"| {r['floor']} | {r['tower']} | {r['F']} | {r['V']} |")
                lines.append("")

    # Section 7: Shear Weight Ratio
    shear_weight = out_data.get("wzq_shear_weight_ratio", [])
    if shear_weight:
        lines.append("## 七、楼层剪重比")
        lines.append("")
        lines.append("| 层号 | 塔号 | X向剪重比 | Y向剪重比 |")
        lines.append("|------|------|----------|----------|")
        for r in shear_weight:
            lines.append(f"| {r['floor']} | {r['tower']} | {r['ratio_x']} | {r['ratio_y']} |")
        lines.append("")

    # Section 8: Floor Displacement
    wdisp_cases = out_data.get("wdisp_cases", [])
    if wdisp_cases:
        lines.append("## 八、层间位移")
        lines.append("")
        for case in wdisp_cases:
            lines.append(f"### 工况 {case['case_num']}: {case['title']}")
            lines.append("")
            if case["floors"]:
                lines.append("| 层号 | 塔号 | 最大位移(mm) | 平均位移(mm) |")
                lines.append("|------|------|------------|------------|")
                for f in case["floors"]:
                    lines.append(f"| {f['floor']} | {f['tower']} | {f['max_disp']} | {f['ave_disp']} |")
            if case["max_drift_summary"]:
                lines.append(f"\n**最大层间位移角**: {case['max_drift_summary']}")
            lines.append("")

    # Section 9: Story Stiffness
    stiff_api = detailed.get("story_stiffness", [])
    stiffness_info = out_data.get("stiffness_info", [])
    if stiff_api or stiffness_info:
        lines.append("## 九、层刚度")
        lines.append("")
        if stiff_api:
            lines.append("| 层号 | RJX(kN/m) | RJY(kN/m) | 比值X | 比值Y |")
            lines.append("|------|-----------|-----------|-------|-------|")
            for s in stiff_api:
                lines.append(
                    f"| {s['floor_index']} | {s['RJX']} | {s['RJY']} | "
                    f"{s['ratio_x']} | {s['ratio_y']} |"
                )
        if stiffness_info:
            lines.append("\n### 刚心、偏心率信息")
            lines.append("")
            lines.append("| 层号 | 刚心X | 刚心Y | 质心X | 质心Y | 偏心率X | 偏心率Y |")
            lines.append("|------|-------|-------|-------|-------|--------|--------|")
            for s in stiffness_info:
                lines.append(
                    f"| {s['floor']} | {s.get('Xstif', '')} | {s.get('Ystif', '')} | "
                    f"{s.get('Xmass', '')} | {s.get('Ymass', '')} | "
                    f"{s.get('Eex', '')} | {s.get('Eey', '')} |"
                )
        lines.append("")

    # Section 10: Overturning
    overturning = out_data.get("overturning", [])
    if overturning:
        lines.append("## 十、抗倾覆验算")
        lines.append("")
        lines.append("| 工况 | 抗倾覆力矩 | 倾覆力矩 | 比值 | 零应力区(%) |")
        lines.append("|------|-----------|---------|------|------------|")
        for r in overturning:
            lines.append(f"| {r['case']} | {r['Mr']} | {r['Mov']} | {r['ratio']} | {r['zero_stress']} |")
        lines.append("")

    # Section 11: Stability
    stability = out_data.get("stability", [])
    stability_conclusion = out_data.get("stability_conclusion", [])
    if stability:
        lines.append("## 十一、结构整体稳定验算")
        lines.append("")
        lines.append("| 层号 | X刚度 | Y刚度 | 层高 | 上部重量 | X刚重比 | Y刚重比 |")
        lines.append("|------|-------|-------|------|---------|--------|--------|")
        for r in stability:
            lines.append(
                f"| {r['floor']} | {r['stiff_x']} | {r['stiff_y']} | {r['height']} | "
                f"{r['upper_weight']} | {r['ratio_x']} | {r['ratio_y']} |"
            )
        lines.append("")
        for c in stability_conclusion:
            lines.append(f"- {c}")
        lines.append("")

    # Section 12: Comfort
    comfort = out_data.get("comfort", [])
    if comfort:
        lines.append("## 十二、结构舒适性验算")
        lines.append("")
        for c in comfort:
            lines.append(f"- {c}")
        lines.append("")

    # Section 13: Shear Capacity
    shear_cap = out_data.get("shear_capacity", [])
    if shear_cap:
        lines.append("## 十三、楼层抗剪承载力")
        lines.append("")
        lines.append("| 层号 | 塔号 | X向承载力 | Y向承载力 | 比值X | 比值Y |")
        lines.append("|------|------|----------|----------|-------|-------|")
        for r in shear_cap:
            lines.append(
                f"| {r['floor']} | {r['tower']} | {r['capacity_x']} | "
                f"{r['capacity_y']} | {r['ratio_x']} | {r['ratio_y']} |"
            )
        lines.append("")

    # Section 14: Column Design (from WPJ)
    wpj_columns = out_data.get("wpj_columns", [])
    if wpj_columns:
        lines.append("## 十四、柱配筋验算结果")
        lines.append("")
        lines.append("| 编号 | 截面B*H(mm) | 轴压比 | 配筋率(%) | 配箍率(%) | 角筋面积 |")
        lines.append("|------|------------|--------|----------|----------|---------|")
        for c in wpj_columns:
            lines.append(
                f"| {c['id']} | {c.get('width', '')}*{c.get('height', '')} | "
                f"{c.get('axial_ratio', '')} | {c.get('reinforce_ratio', '')} | "
                f"{c.get('hoop_ratio', '')} | {c.get('corner_steel', '')} |"
            )
        lines.append("")

    # Section 15: Beam Design (from WPJ)
    wpj_beams = out_data.get("wpj_beams", [])
    if wpj_beams:
        lines.append("## 十五、梁配筋验算结果")
        lines.append("")
        lines.append("| 编号 | 截面B*H(mm) | 上部配筋 | 下部配筋 |")
        lines.append("|------|------------|---------|---------|")
        for b in wpj_beams:
            lines.append(
                f"| {b['id']} | {b.get('width', '')}*{b.get('height', '')} | "
                f"{b.get('top_reinforce', '')} | {b.get('btm_reinforce', '')} |"
            )
        lines.append("")

    # Section 16: Exceedance
    wgcpj = out_data.get("wgcpj", {})
    exceed_items = wgcpj.get("items", [])
    if exceed_items:
        lines.append("## 十六、超限信息")
        lines.append("")
        for item in exceed_items:
            lines.append(f"- {item}")
        lines.append("")

    # Section 17: Beam/Column Summary (API data)
    beams = detailed.get("beam_design", {})
    if beams.get("total_beams", 0) > 0:
        lines.append("## 十七、梁设计统计")
        lines.append("")
        lines.append(f"- **总梁数**: {beams['total_beams']}")
        lines.append(f"- **最大剪压比**: {beams['max_shear_compression_ratio']}")
        lines.append(f"- **总配筋量**: {beams['total_reinforce_quantity']}")
        lines.append("")

    cols = detailed.get("column_design", {})
    if cols.get("total_columns", 0) > 0:
        lines.append("## 十八、柱设计统计")
        lines.append("")
        lines.append(f"- **总柱数**: {cols['total_columns']}")
        lines.append(f"- **最大轴压比**: {cols['max_axial_compression_ratio']}")
        lines.append(f"- **总配筋量**: {cols['total_reinforce_quantity']}")
        lines.append("")

    return "\n".join(lines)


# ── Word document generation ────────────────────────────────────────────


def _add_table(doc: Any, headers: List[str], rows: List[List[str]]) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Pt

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
        shading = cell._element.get_or_add_tcPr()
        shading_elem = shading.makeelement(qn("w:shd"), {
            qn("w:fill"): "D9E2F3",
            qn("w:val"): "clear",
        })
        shading.append(shading_elem)

    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(10)


def _add_kv_section(doc: Any, title: str, text: str) -> None:
    doc.add_heading(title, level=2)
    for line in text.strip().split("\n"):
        line = line.strip()
        if line:
            doc.add_paragraph(line, style="List Bullet")


def _generate_docx(report: Dict[str, Any], output_path: Path, images: Optional[List[Path]] = None) -> Path:
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx import Document

    doc = Document()
    title = doc.add_heading("PKPM SATWE 结构计算书", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(22)

    # Insert structural images after title
    if images:
        doc.add_heading("结构平面图", level=1)
        for img_path in images:
            embed_path = _convert_bmp_to_png(img_path) if img_path.suffix.lower() == ".bmp" else img_path
            if embed_path is None:
                _logger.warning("Skipping unembeddable image: %s", img_path.name)
                continue
            try:
                doc.add_picture(str(embed_path), width=Inches(6))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as exc:
                _logger.warning("Failed to embed image %s in DOCX: %s", img_path.name, exc)

    detailed = report.get("detailed", {})
    out_data = detailed.get("out_file_data", {})
    params = out_data.get("wmass_params", {})

    # Section 1
    if params:
        doc.add_heading("一、设计参数总信息", level=1)
        for key, label in [
            ("design_params", "总体信息"),
            ("wind_info_params", "风荷载参数"),
            ("earthquake_params", "地震参数"),
            ("material_params", "材料参数"),
        ]:
            section = params.get(key)
            if section:
                _add_kv_section(doc, label, section)

    # Section 2: Mass
    mass_table = out_data.get("mass_table", [])
    total_mass = out_data.get("total_mass", {})
    member_counts = out_data.get("member_counts", [])
    if mass_table or member_counts:
        doc.add_heading("二、楼层质量与构件信息", level=1)
        if total_mass:
            for k, v in total_mass.items():
                doc.add_paragraph(f"{k}: {v} t")
        if mass_table:
            doc.add_heading("各层质量", level=2)
            _add_table(doc,
                ["层号", "塔号", "恒载质量(t)", "活载质量(t)", "质量比"],
                [[str(r["floor"]), str(r["tower"]), str(r["dead_mass"]),
                  str(r["live_mass"]), str(r["mass_ratio"])] for r in mass_table])
        if member_counts:
            doc.add_heading("构件数量与层高", level=2)
            _add_table(doc,
                ["层号", "梁数", "柱数", "层高(m)", "累计高度(m)"],
                [[str(r["floor"]), str(r["beam_count"]), str(r["column_count"]),
                  str(r["height"]), str(r["cumulative_height"])] for r in member_counts])

    # Section 3: Unit Mass
    unit_mass = out_data.get("unit_mass", [])
    if unit_mass:
        doc.add_heading("三、单位面积质量分布", level=1)
        _add_table(doc,
            ["层号", "塔号", "单位面积质量(kg/m²)", "质量比"],
            [[str(r["floor"]), str(r["tower"]), str(r["unit_mass"]), str(r["mass_ratio"])]
             for r in unit_mass])

    # Section 4: Wind Load
    wind_load = out_data.get("wind_load", [])
    if wind_load:
        doc.add_heading("四、风荷载信息", level=1)
        _add_table(doc,
            ["层号", "风荷载X", "剪力X", "倾覆X", "风荷载Y", "剪力Y", "倾覆Y"],
            [[str(r["floor"]), str(r["wind_x"]), str(r["shear_x"]), str(r["overturn_x"]),
              str(r["wind_y"]), str(r["shear_y"]), str(r["overturn_y"])] for r in wind_load])

    # Section 5: Modal
    wzq_periods = out_data.get("wzq_periods", [])
    modal_api = detailed.get("modal_analysis", [])
    if wzq_periods:
        doc.add_heading("五、模态分析", level=1)
        _add_table(doc,
            ["振型号", "周期(s)", "方向角(°)", "平动系数", "X平动", "Y平动", "扭转系数"],
            [[str(m["mode"]), str(m["period"]), str(m["angle"]), str(m["translation"]),
              str(m["x_translation"]), str(m["y_translation"]), str(m["torsion"])]
             for m in wzq_periods])
    elif modal_api:
        doc.add_heading("五、模态分析", level=1)
        _add_table(doc,
            ["振型号", "周期(s)", "方向角(°)", "扭转系数", "X侧移", "Y侧移"],
            [[str(m["index"]), str(m["period_s"]), str(m["angle"]),
              str(m["torsion_ratio"]), str(m["x_side"]), str(m["y_side"])]
             for m in modal_api])

    # Section 6: Earthquake CQC
    wzq_base_shear = out_data.get("wzq_base_shear", [])
    if wzq_base_shear:
        doc.add_heading("六、地震作用下的楼层反应（CQC）", level=1)
        for direction in ["X", "Y"]:
            dir_data = [r for r in wzq_base_shear if r["direction"] == direction]
            if dir_data:
                doc.add_heading(f"{direction}方向", level=2)
                _add_table(doc,
                    ["层号", "塔号", "F(kN)", "V(kN)"],
                    [[str(r["floor"]), str(r["tower"]), str(r["F"]), str(r["V"])]
                     for r in dir_data])

    # Section 7: Shear Weight Ratio
    shear_weight = out_data.get("wzq_shear_weight_ratio", [])
    if shear_weight:
        doc.add_heading("七、楼层剪重比", level=1)
        _add_table(doc,
            ["层号", "塔号", "X向剪重比", "Y向剪重比"],
            [[str(r["floor"]), str(r["tower"]), str(r["ratio_x"]), str(r["ratio_y"])]
             for r in shear_weight])

    # Section 8: Displacement
    wdisp_cases = out_data.get("wdisp_cases", [])
    if wdisp_cases:
        doc.add_heading("八、层间位移", level=1)
        for case in wdisp_cases:
            doc.add_heading(f"工况 {case['case_num']}: {case['title']}", level=2)
            if case["floors"]:
                _add_table(doc,
                    ["层号", "塔号", "最大位移(mm)", "平均位移(mm)"],
                    [[str(f["floor"]), str(f["tower"]), str(f["max_disp"]), str(f["ave_disp"])]
                     for f in case["floors"]])
            if case["max_drift_summary"]:
                doc.add_paragraph(f"最大层间位移角: {case['max_drift_summary']}")

    # Section 9: Stiffness
    stiff_api = detailed.get("story_stiffness", [])
    if stiff_api:
        doc.add_heading("九、层刚度", level=1)
        _add_table(doc,
            ["层号", "RJX(kN/m)", "RJY(kN/m)", "比值X", "比值Y"],
            [[str(s["floor_index"]), str(s["RJX"]), str(s["RJY"]),
              str(s["ratio_x"]), str(s["ratio_y"])] for s in stiff_api])

    # Section 10: Overturning
    overturning = out_data.get("overturning", [])
    if overturning:
        doc.add_heading("十、抗倾覆验算", level=1)
        _add_table(doc,
            ["工况", "抗倾覆力矩", "倾覆力矩", "比值", "零应力区(%)"],
            [[str(r["case"]), str(r["Mr"]), str(r["Mov"]), str(r["ratio"]), str(r["zero_stress"])]
             for r in overturning])

    # Section 11: Stability
    stability = out_data.get("stability", [])
    if stability:
        doc.add_heading("十一、结构整体稳定验算", level=1)
        _add_table(doc,
            ["层号", "X刚度", "Y刚度", "层高", "上部重量", "X刚重比", "Y刚重比"],
            [[str(r["floor"]), str(r["stiff_x"]), str(r["stiff_y"]), str(r["height"]),
              str(r["upper_weight"]), str(r["ratio_x"]), str(r["ratio_y"])]
             for r in stability])
        for c in out_data.get("stability_conclusion", []):
            doc.add_paragraph(c)

    # Section 12: Comfort
    comfort = out_data.get("comfort", [])
    if comfort:
        doc.add_heading("十二、结构舒适性验算", level=1)
        for c in comfort:
            doc.add_paragraph(c, style="List Bullet")

    # Section 13: Shear Capacity
    shear_cap = out_data.get("shear_capacity", [])
    if shear_cap:
        doc.add_heading("十三、楼层抗剪承载力", level=1)
        _add_table(doc,
            ["层号", "塔号", "X向承载力", "Y向承载力", "比值X", "比值Y"],
            [[str(r["floor"]), str(r["tower"]), str(r["capacity_x"]),
              str(r["capacity_y"]), str(r["ratio_x"]), str(r["ratio_y"])]
             for r in shear_cap])

    # Section 14: Column Design
    wpj_columns = out_data.get("wpj_columns", [])
    if wpj_columns:
        doc.add_heading("十四、柱配筋验算结果", level=1)
        _add_table(doc,
            ["编号", "截面B*H(mm)", "轴压比", "配筋率(%)", "配箍率(%)", "角筋面积"],
            [[str(c["id"]), f"{c.get('width', '')}*{c.get('height', '')}",
              str(c.get("axial_ratio", "")), str(c.get("reinforce_ratio", "")),
              str(c.get("hoop_ratio", "")), str(c.get("corner_steel", ""))]
             for c in wpj_columns])

    # Section 15: Beam Design
    wpj_beams = out_data.get("wpj_beams", [])
    if wpj_beams:
        doc.add_heading("十五、梁配筋验算结果", level=1)
        _add_table(doc,
            ["编号", "截面B*H(mm)", "上部配筋", "下部配筋"],
            [[str(b["id"]), f"{b.get('width', '')}*{b.get('height', '')}",
              str(b.get("top_reinforce", "")), str(b.get("btm_reinforce", ""))]
             for b in wpj_beams])

    # Section 16: Exceedance
    wgcpj = out_data.get("wgcpj", {})
    exceed_items = wgcpj.get("items", [])
    if exceed_items:
        doc.add_heading("十六、超限信息", level=1)
        for item in exceed_items:
            doc.add_paragraph(item, style="List Bullet")

    # Section 17-18: API Summary
    beams = detailed.get("beam_design", {})
    if beams.get("total_beams", 0) > 0:
        doc.add_heading("十七、梁设计统计", level=1)
        doc.add_paragraph(f"总梁数: {beams['total_beams']}")
        doc.add_paragraph(f"最大剪压比: {beams['max_shear_compression_ratio']}")
        doc.add_paragraph(f"总配筋量: {beams['total_reinforce_quantity']}")

    cols = detailed.get("column_design", {})
    if cols.get("total_columns", 0) > 0:
        doc.add_heading("十八、柱设计统计", level=1)
        doc.add_paragraph(f"总柱数: {cols['total_columns']}")
        doc.add_paragraph(f"最大轴压比: {cols['max_axial_compression_ratio']}")
        doc.add_paragraph(f"总配筋量: {cols['total_reinforce_quantity']}")

    doc.save(str(output_path))
    return output_path


# ── PDF generation ──────────────────────────────────────────────────────


def _generate_pdf(report: Dict[str, Any], output_path: Path, images: Optional[List[Path]] = None) -> Path:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.platypus import Image as RLImage, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    cn_font = "STSong-Light"

    doc = SimpleDocTemplate(
        str(output_path), pagesize=A4,
        leftMargin=20 * mm, rightMargin=20 * mm,
        topMargin=20 * mm, bottomMargin=20 * mm,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("CNTitle", parent=styles["Title"], fontName=cn_font, fontSize=18, spaceAfter=12)
    h1_style = ParagraphStyle("CNH1", parent=styles["Heading1"], fontName=cn_font, fontSize=14, spaceAfter=8, spaceBefore=12)
    h2_style = ParagraphStyle("CNH2", parent=styles["Heading2"], fontName=cn_font, fontSize=12, spaceAfter=6, spaceBefore=8)
    body_style = ParagraphStyle("CNBody", parent=styles["Normal"], fontName=cn_font, fontSize=10, spaceAfter=4)

    elements: List[Any] = []
    detailed = report.get("detailed", {})
    out_data = detailed.get("out_file_data", {})

    elements.append(Paragraph("PKPM SATWE 结构计算书", title_style))
    elements.append(Spacer(1, 10))

    # Insert structural images after title
    if images:
        elements.append(Paragraph("结构平面图", h1_style))
        for img_path in images:
            embed_path = _convert_bmp_to_png(img_path) if img_path.suffix.lower() == ".bmp" else img_path
            if embed_path is None:
                _logger.warning("Skipping unembeddable image: %s", img_path.name)
                continue
            try:
                avail_w = doc.width
                img = RLImage(str(embed_path), width=avail_w, height=avail_w * 0.5)
                img.hAlign = "CENTER"
                elements.append(img)
                elements.append(Spacer(1, 8))
            except Exception as exc:
                _logger.warning("Failed to embed image %s in PDF: %s", img_path.name, exc)

    def _add_pdf_table(headers: List[str], rows: List[List[str]]) -> None:
        data = [headers] + rows
        col_count = len(headers)
        avail = doc.width
        col_w = [avail / col_count] * col_count
        t = Table(data, colWidths=col_w)
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#D9E2F3")),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("FONTNAME", (0, 0), (-1, -1), cn_font),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F2F2F2")]),
        ]))
        elements.append(t)
        elements.append(Spacer(1, 8))

    # Section 1: Design Params
    params = out_data.get("wmass_params", {})
    if params:
        elements.append(Paragraph("一、设计参数总信息", h1_style))
        for key, label in [
            ("design_params", "总体信息"),
            ("wind_info_params", "风荷载参数"),
            ("earthquake_params", "地震参数"),
            ("material_params", "材料参数"),
        ]:
            section = params.get(key)
            if section:
                elements.append(Paragraph(label, h2_style))
                for line in section.strip().split("\n"):
                    line = line.strip()
                    if line:
                        elements.append(Paragraph(line, body_style))

    # Section 2: Mass
    mass_table = out_data.get("mass_table", [])
    if mass_table:
        elements.append(Paragraph("二、各层质量", h1_style))
        _add_pdf_table(
            ["层号", "塔号", "恒载(t)", "活载(t)", "质量比"],
            [[str(r["floor"]), str(r["tower"]), str(r["dead_mass"]),
              str(r["live_mass"]), str(r["mass_ratio"])] for r in mass_table])

    # Section 3: Unit Mass
    unit_mass = out_data.get("unit_mass", [])
    if unit_mass:
        elements.append(Paragraph("三、单位面积质量分布", h1_style))
        _add_pdf_table(
            ["层号", "塔号", "单位质量(kg/m²)", "质量比"],
            [[str(r["floor"]), str(r["tower"]), str(r["unit_mass"]), str(r["mass_ratio"])]
             for r in unit_mass])

    # Section 4: Wind
    wind_load = out_data.get("wind_load", [])
    if wind_load:
        elements.append(Paragraph("四、风荷载信息", h1_style))
        _add_pdf_table(
            ["层号", "风X", "剪力X", "倾覆X", "风Y", "剪力Y", "倾覆Y"],
            [[str(r["floor"]), str(r["wind_x"]), str(r["shear_x"]), str(r["overturn_x"]),
              str(r["wind_y"]), str(r["shear_y"]), str(r["overturn_y"])] for r in wind_load])

    # Section 5: Modal
    wzq_periods = out_data.get("wzq_periods", [])
    modal_api = detailed.get("modal_analysis", [])
    if wzq_periods:
        elements.append(Paragraph("五、模态分析", h1_style))
        _add_pdf_table(
            ["振型号", "周期(s)", "方向角", "平动", "扭转"],
            [[str(m["mode"]), str(m["period"]), str(m["angle"]),
              str(m["translation"]), str(m["torsion"])] for m in wzq_periods])
    elif modal_api:
        elements.append(Paragraph("五、模态分析", h1_style))
        _add_pdf_table(
            ["振型号", "周期(s)", "方向角", "扭转系数", "X侧移", "Y侧移"],
            [[str(m["index"]), str(m["period_s"]), str(m["angle"]),
              str(m["torsion_ratio"]), str(m["x_side"]), str(m["y_side"])]
             for m in modal_api])

    # Section 6: CQC
    wzq_base_shear = out_data.get("wzq_base_shear", [])
    if wzq_base_shear:
        elements.append(Paragraph("六、地震楼层反应(CQC)", h1_style))
        for direction in ["X", "Y"]:
            dir_data = [r for r in wzq_base_shear if r["direction"] == direction]
            if dir_data:
                elements.append(Paragraph(f"{direction}方向", h2_style))
                _add_pdf_table(
                    ["层号", "塔号", "F(kN)", "V(kN)"],
                    [[str(r["floor"]), str(r["tower"]), str(r["F"]), str(r["V"])]
                     for r in dir_data])

    # Section 7: Shear Weight
    shear_weight = out_data.get("wzq_shear_weight_ratio", [])
    if shear_weight:
        elements.append(Paragraph("七、楼层剪重比", h1_style))
        _add_pdf_table(
            ["层号", "塔号", "X向", "Y向"],
            [[str(r["floor"]), str(r["tower"]), str(r["ratio_x"]), str(r["ratio_y"])]
             for r in shear_weight])

    # Section 8: Displacement
    wdisp_cases = out_data.get("wdisp_cases", [])
    if wdisp_cases:
        elements.append(Paragraph("八、层间位移", h1_style))
        for case in wdisp_cases:
            elements.append(Paragraph(f"工况{case['case_num']}: {case['title']}", h2_style))
            if case["floors"]:
                _add_pdf_table(
                    ["层号", "塔号", "最大位移", "平均位移"],
                    [[str(f["floor"]), str(f["tower"]), str(f["max_disp"]), str(f["ave_disp"])]
                     for f in case["floors"]])

    # Section 9: Stiffness
    stiff_api = detailed.get("story_stiffness", [])
    if stiff_api:
        elements.append(Paragraph("九、层刚度", h1_style))
        _add_pdf_table(
            ["层号", "RJX(kN/m)", "RJY(kN/m)", "比值X", "比值Y"],
            [[str(s["floor_index"]), str(s["RJX"]), str(s["RJY"]),
              str(s["ratio_x"]), str(s["ratio_y"])] for s in stiff_api])

    # Section 10: Overturning
    overturning = out_data.get("overturning", [])
    if overturning:
        elements.append(Paragraph("十、抗倾覆验算", h1_style))
        _add_pdf_table(
            ["工况", "抗倾覆力矩", "倾覆力矩", "比值", "零应力区(%)"],
            [[str(r["case"]), str(r["Mr"]), str(r["Mov"]), str(r["ratio"]), str(r["zero_stress"])]
             for r in overturning])

    # Section 11: Stability
    stability = out_data.get("stability", [])
    if stability:
        elements.append(Paragraph("十一、结构整体稳定验算", h1_style))
        _add_pdf_table(
            ["层号", "X刚度", "Y刚度", "层高", "X刚重比", "Y刚重比"],
            [[str(r["floor"]), str(r["stiff_x"]), str(r["stiff_y"]), str(r["height"]),
              str(r["ratio_x"]), str(r["ratio_y"])] for r in stability])

    # Section 12: Column Design
    wpj_columns = out_data.get("wpj_columns", [])
    if wpj_columns:
        elements.append(Paragraph("十四、柱配筋验算结果", h1_style))
        _add_pdf_table(
            ["编号", "截面", "轴压比", "配筋率(%)", "配箍率(%)"],
            [[str(c["id"]), f"{c.get('width', '')}*{c.get('height', '')}",
              str(c.get("axial_ratio", "")), str(c.get("reinforce_ratio", "")),
              str(c.get("hoop_ratio", ""))] for c in wpj_columns])

    # Section 13: Beam Design
    wpj_beams = out_data.get("wpj_beams", [])
    if wpj_beams:
        elements.append(Paragraph("十五、梁配筋验算结果", h1_style))
        _add_pdf_table(
            ["编号", "截面", "上部配筋", "下部配筋"],
            [[str(b["id"]), f"{b.get('width', '')}*{b.get('height', '')}",
              str(b.get("top_reinforce", "")), str(b.get("btm_reinforce", ""))]
             for b in wpj_beams])

    # Section: Exceedance
    wgcpj = out_data.get("wgcpj", {})
    exceed_items = wgcpj.get("items", [])
    if exceed_items:
        elements.append(Paragraph("十六、超限信息", h1_style))
        for item in exceed_items:
            elements.append(Paragraph(item, body_style))

    # Section: API Summary
    beams = detailed.get("beam_design", {})
    if beams.get("total_beams", 0) > 0:
        elements.append(Paragraph("十七、梁设计统计", h1_style))
        elements.append(Paragraph(f"总梁数: {beams['total_beams']}", body_style))
        elements.append(Paragraph(f"最大剪压比: {beams['max_shear_compression_ratio']}", body_style))

    cols = detailed.get("column_design", {})
    if cols.get("total_columns", 0) > 0:
        elements.append(Paragraph("十八、柱设计统计", h1_style))
        elements.append(Paragraph(f"总柱数: {cols['total_columns']}", body_style))
        elements.append(Paragraph(f"最大轴压比: {cols['max_axial_compression_ratio']}", body_style))

    doc.build(elements)
    return output_path


# ── PDF conversion fallback ──────────────────────────────────────────────


def _find_wps_exe() -> Optional[Path]:
    base = Path(os.environ.get("LOCALAPPDATA", "")) / "Kingsoft" / "WPS Office"
    if not base.is_dir():
        return None
    for ver_dir in sorted(base.iterdir(), reverse=True):
        candidate = ver_dir / "office6" / "wps.exe"
        if candidate.is_file():
            return candidate
    return None


def _convert_docx_to_pdf(docx_path: Path) -> Optional[Path]:
    pdf_path = docx_path.with_suffix(".pdf")
    wps_exe = _find_wps_exe()
    if wps_exe:
        try:
            subprocess.run(
                [str(wps_exe), str(docx_path), "/ExportPDF", str(pdf_path)],
                stdout=subprocess.PIPE, stderr=subprocess.PIPE,
                timeout=30,
            )
            if pdf_path.is_file():
                return pdf_path
        except Exception as exc:
            _logger.debug("WPS CLI conversion failed: %s", exc)
    try:
        import win32com.client
        import pythoncom
        pythoncom.CoInitialize()
        try:
            wps = win32com.client.Dispatch("kwps.Application")
            wps.Visible = False
            doc = wps.Documents.Open(str(docx_path.resolve()))
            doc.ExportAsFixedFormat(str(pdf_path.resolve()), 0)
            doc.Close(False)
            wps.Quit()
        finally:
            pythoncom.CoUninitialize()
        if pdf_path.is_file():
            return pdf_path
    except Exception as exc:
        _logger.debug("COM-based PDF conversion failed: %s", exc)
    return None
